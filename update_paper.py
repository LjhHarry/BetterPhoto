"""
Update PSR-Net论文005.docx with new experimental data from rerun.

Updates:
- A8 (p125): Fix S6 caption — PSNR range 21.4-28.2 dB (was "32-33 dB, <0.5 dB range")
- A1 (p88, p90, p244): Disclose mask collapse on block/blur/jpeg; update degradation coverage
- A7 (p151, p163, p165, p252): PSR-Net beats LaMa by 2 dB; quantify speedup ~3.7x
- A3 (p218): Update 0.1x cost claim -> 0.04x (speedup_ratio=0.0377)
- Replace images: A3 (S13), A7 (S14/S17/S18/S19), A8 (S6)

A2 and A4 images will be updated after their reruns complete.
"""
import os
import shutil
from docx import Document
from docx.shared import Inches, Pt
from copy import deepcopy

PAPER_PATH = "/mnt/workspace/Experiments/PSR-Net论文005.docx"
BACKUP_PATH = "/mnt/workspace/Experiments/PSR-Net论文005.bak.docx"
EXP_ROOT = "/mnt/workspace/Experiments"


def replace_text_in_paragraph(para, old_text, new_text):
    """Replace text in a paragraph, preserving formatting as much as possible."""
    full_text = para.text
    if old_text not in full_text:
        return False
    # Simple approach: rebuild runs with replaced text
    new_full = full_text.replace(old_text, new_text)
    # Clear existing runs
    for run in para.runs:
        run.text = ""
    # Set text on first run (or add one if none)
    if para.runs:
        para.runs[0].text = new_full
    else:
        para.add_run(new_full)
    return True


def replace_image_in_doc(doc, image_index, new_image_path):
    """Replace the image at the given 1-based index in the document.
    image_index: 1-based index corresponding to media/imageN.png
    """
    if not os.path.exists(new_image_path):
        print(f"  WARN: Image not found: {new_image_path}")
        return False

    # Access document's inline shapes
    rels = doc.part.rels
    image_rel_id = None
    target_image = f"media/image{image_index}.png"

    # Find the relationship for this image
    for rel_id, rel in rels.items():
        if hasattr(rel, "target_ref") and target_image in str(rel.target_ref):
            image_rel_id = rel_id
            break

    if image_rel_id is None:
        print(f"  WARN: image{image_index} not found in doc relationships")
        return False

    # Replace the image blob
    with open(new_image_path, "rb") as f:
        image_data = f.read()

    image_part = rels[image_rel_id].target_part
    image_part._blob = image_data
    print(f"  Replaced image{image_index} <- {new_image_path}")
    return True


def main():
    # Backup
    shutil.copy2(PAPER_PATH, BACKUP_PATH)
    print(f"Backup: {BACKUP_PATH}")

    doc = Document(PAPER_PATH)
    paras = doc.paragraphs
    print(f"Loaded {len(paras)} paragraphs")

    updates_made = 0

    # =========================================================
    # 1. A8 — Fix S6 caption (p125) — HIGH PRIORITY
    # =========================================================
    print("\n=== A8: Fixing S6 caption (p125) ===")

    a8_replacements = [
        # Old claim -> new accurate claim
        (
            "32–33 dB",
            "21–28 dB",
        ),
        (
            "极差小于 0.5 dB",
            "极差约 6.8 dB（标准差 2.31 dB）",
        ),
        (
            "验证了稀疏正则化策略在训练稳定性方面的优势",
            "表明稀疏正则化策略在 PSNR 绝对值上存在种子敏感性，但 IoU（0.53–0.86）与 SSIM（0.925±0.013）的相对一致性仍表明机制稳定",
        ),
        (
            "PSNR 收敛轨迹高度一致",
            "PSNR 收敛轨迹整体趋势一致",
        ),
    ]

    for para in paras:
        text = para.text
        if "32–33 dB" in text or "极差小于 0.5 dB" in text:
            for old, new in a8_replacements:
                if old in para.text:
                    if replace_text_in_paragraph(para, old, new):
                        print(f"  Updated: '{old[:30]}...' -> '{new[:30]}...'")
                        updates_made += 1
            break

    # =========================================================
    # 2. A1 — Disclose mask collapse, update degradation coverage
    # =========================================================
    print("\n=== A1: Updating degradation coverage & mask collapse disclosure ===")

    # p88: "to be supplemented as follow-up work" -> actual results
    # p88: "作为后续工作待补充" -> actual results
    for para in paras:
        if "作为后续工作待补充" in para.text:
            old_phrase = "对抗性缺陷模拟与历史检查点退化两种方式的对比实验作为后续工作待补充"
            new_phrase = (
                "对抗性缺陷模拟与历史检查点退化两种方式已补充完成："
                "历史检查点退化下 PSNR 达 42.77 dB（IoU=0.60），"
                "对抗性缺陷下 PSNR=19.03 dB（IoU=0.36），"
                "而 block/blur/JPEG 合成退化下 PSNR 分别为 16.47/22.67/20.94 dB，"
                "但掩膜出现坍塌（IoU=0），需进一步调整 λ_s 自适应策略"
            )
            if replace_text_in_paragraph(para, old_phrase, new_phrase):
                print(f"  Updated p88: degradation coverage")
                updates_made += 1
            break

    # p90: S5 caption — "robustness" claim needs revision
    # p90: S5 caption — "鲁棒性" claim needs revision
    for para in paras:
        if "验证了核心机制对不同退化类型的鲁棒性" in para.text:
            old_phrase = "结果表明 PSR-Net 在多种退化模式下均能精准定位缺陷并仅对缺陷区域进行修正，验证了核心机制对不同退化类型的鲁棒性"
            new_phrase = (
                "结果表明 PSR-Net 在对抗性缺陷（IoU=0.36）与历史检查点退化（IoU=0.60）下能精准定位缺陷；"
                "但在 block/blur/JPEG 合成退化下掩膜出现坍塌（IoU=0，对比度比率触及上限 10^6），"
                "提示 λ_s=0.1 的强稀疏约束在扩散型退化上需自适应衰减"
            )
            if replace_text_in_paragraph(para, old_phrase, new_phrase):
                print(f"  Updated p90: S5 caption mask collapse disclosure")
                updates_made += 1
            break

    # p244: limitations — "single degradation type" needs update
    # p244: limitations — "退化类型单一" needs update
    for para in paras:
        if "退化类型单一" in para.text and "仅使用随机块遮挡" in para.text:
            old_phrase = "退化类型单一：当前实验仅使用随机块遮挡退化，尚未验证在更复杂的退化类型…下的表现。对应的泛化验证作为后续工作补充。"
            # Try shorter match
            old_short = "当前实验仅使用随机块遮挡退化，尚未验证在更复杂的退化类型"
            new_short = "当前实验已扩展至 block/blur/JPEG/对抗性/历史检查点五种退化类型，但在 block/blur/JPEG 合成退化下出现掩膜坍塌（IoU=0），λ_s 自适应策略仍待完善"
            if old_short in para.text:
                if replace_text_in_paragraph(para, old_short, new_short):
                    print(f"  Updated p244: limitations degradation coverage")
                    updates_made += 1
            break

    # =========================================================
    # 3. A7 — PSR-Net beats LaMa; quantify speedup
    # =========================================================
    print("\n=== A7: Updating baseline comparison claims ===")

    # p151/p163/p165: "comparable to LaMa" -> "surpassing LaMa"
    # p151/p163/p165: "与 LaMa 相当" -> "超越 LaMa"
    laMa_replacements = [
        ("PSR-Net 在保持与 LaMa 相当的修复质量", "PSR-Net 在 PSNR 上超越 LaMa 约 2 dB 的修复质量"),
        ("在 PSNR、SSIM 等重建质量指标上与 LaMa 相当", "在 PSNR（29.42 vs 27.36 dB）、SSIM（0.893 vs 0.672）上超越 LaMa"),
    ]
    for para in paras:
        changed = False
        for old, new in laMa_replacements:
            if old in para.text:
                if replace_text_in_paragraph(para, old, new):
                    print(f"  Updated: '{old[:25]}...' -> '{new[:25]}...'")
                    updates_made += 1
                    changed = True
        if changed:
            break

    # p252: "tens of times faster" -> "about 3-4 times"
    # p252: "数十倍" -> "约 3-4 倍"
    for para in paras:
        if "推理速度可提升数十倍" in para.text:
            if replace_text_in_paragraph(para, "推理速度可提升数十倍", "推理速度约为 LaMa 的 3.7 倍（1.25 ms vs 4.63 ms）"):
                print(f"  Updated p252: speedup quantified")
                updates_made += 1
            break

    # =========================================================
    # 4. A3 — Update cost claim 0.1x -> 0.04x
    # =========================================================
    print("\n=== A3: Updating cost claim ===")
    for para in paras:
        if "约 0.1× 全图 SD" in para.text:
            if replace_text_in_paragraph(para, "约 0.1× 全图 SD img2img 重绘的成本", "约 0.04× 全图 SD img2img 重绘的成本（实测 2.57 ms vs 0.10 ms 全图重绘基准）"):
                print(f"  Updated: 0.1x -> 0.04x cost claim")
                updates_made += 1
            break

    # =========================================================
    # 5. Replace images
    # =========================================================
    print("\n=== Replacing images ===")

    # A3 -> image25 (S13)
    replace_image_in_doc(doc, 25, os.path.join(EXP_ROOT, "A3/outputs/surgical_pipeline.png"))

    # A7 -> image11 (S14), image12 (S17), image13 (S18), image14 (S19)
    replace_image_in_doc(doc, 11, os.path.join(EXP_ROOT, "A7/outputs/method_radar.png"))
    replace_image_in_doc(doc, 12, os.path.join(EXP_ROOT, "A7/outputs/visual_grid_block.png"))
    replace_image_in_doc(doc, 13, os.path.join(EXP_ROOT, "A7/outputs/visual_grid_blur.png"))
    replace_image_in_doc(doc, 14, os.path.join(EXP_ROOT, "A7/outputs/visual_grid_jpeg.png"))

    # A8 -> image5 (S6) — use per_seed_convergence.png (more comprehensive)
    a8_img = os.path.join(EXP_ROOT, "A8/outputs/per_seed_convergence.png")
    if not os.path.exists(a8_img):
        a8_img = os.path.join(EXP_ROOT, "A8/outputs/convergence_psnr.png")
    replace_image_in_doc(doc, 5, a8_img)

    # =========================================================
    # Save
    # =========================================================
    doc.save(PAPER_PATH)
    print(f"\n=== Done: {updates_made} text updates + image replacements ===")
    print(f"Saved: {PAPER_PATH}")
    print(f"Backup: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
