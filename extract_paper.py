"""Extract structure of PSR-Net paper for update planning."""
import json
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

PAPER = Path("/mnt/workspace/Experiments/PSR-Net论文005.docx")
OUT = Path("/mnt/workspace/Experiments/paper_structure.json")

doc = Document(PAPER)

# Build a body-order list of paragraphs and tables by walking the body XML.
# Each item in doc.element.body is either a <w:p> (paragraph) or <w:tbl> (table).
body = doc.element.body
para_idx = 0
table_idx = 0
items = []  # list of dicts: {kind: 'para'|'table', idx, ...}

# Pre-collect paragraphs and tables in document order for indexing
para_map = {}  # id(p_xml) -> index in doc.paragraphs
for i, p in enumerate(doc.paragraphs):
    para_map[id(p._p)] = i

table_map = {}
for i, t in enumerate(doc.tables):
    table_map[id(t._tbl)] = i

# Walk the body in order
for child in body.iterchildren():
    tag = child.tag
    if tag == qn("w:p"):
        p_idx = para_map.get(id(child))
        if p_idx is not None:
            p = doc.paragraphs[p_idx]
            items.append({
                "kind": "para",
                "idx": p_idx,
                "text": p.text,
                "style": p.style.name if p.style else None,
            })
            para_idx = p_idx + 1
    elif tag == qn("w:tbl"):
        t_idx = table_map.get(id(child))
        if t_idx is not None:
            t = doc.tables[t_idx]
            rows = []
            for r in t.rows:
                rows.append([c.text for c in r.cells])
            items.append({
                "kind": "table",
                "idx": t_idx,
                "rows": rows,
                "n_rows": len(rows),
                "n_cols": len(rows[0]) if rows else 0,
            })

# Find images: map drawing elements to paragraph index
image_locations = []  # {para_idx, embed_rId, image_filename}
rels = doc.part.rels
# Build rId -> filename map
rid_to_target = {rid: r.target_ref for rid, r in rels.items()}

for i, p in enumerate(doc.paragraphs):
    # find blip elements
    blips = p._p.findall(".//" + qn("a:blip"))
    for b in blips:
        embed = b.get(qn("r:embed"))
        if embed:
            target = rid_to_target.get(embed, "?")
            image_locations.append({
                "para_idx": i,
                "rid": embed,
                "target": target,
                "paragraph_text": p.text,
            })

# Find all numeric metric references in paragraphs
# Patterns: "32.9 dB", "382527×", "0.87", "PSNR=29.05", percentages, etc.
metric_patterns = [
    (r"\d+\.\d+\s*dB", "psnr_db"),
    (r"\d+\.\d+\s*\\times", "times_latex"),
    (r"\d+\.\d+\s*×", "times_unicode"),
    (r"\d{4,}\s*×", "big_times_unicode"),
    (r"\d+\.\d+±\d+\.\d+", "mean_std"),
    (r"IoU\s*[=为是]\s*0?\.\d+", "iou"),
    (r"PSNR\s*[=为是]\s*\d+\.\d+", "psnr_assign"),
    (r"SSIM\s*[=为是]\s*0?\.\d+", "ssim_assign"),
    (r"F1\s*[=为是]\s*0?\.\d+", "f1_assign"),
]

metric_hits = []
for i, p in enumerate(doc.paragraphs):
    text = p.text
    if not text.strip():
        continue
    for pat, label in metric_patterns:
        for m in re.finditer(pat, text):
            metric_hits.append({
                "para_idx": i,
                "type": label,
                "match": m.group(),
                "context": text[max(0, m.start()-30):min(len(text), m.end()+30)],
            })

# Section outline: paragraphs styled Heading 1-4
outline = []
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else ""
    if style.startswith("Heading") or style.startswith("标题") or style in ("Title",):
        outline.append({
            "para_idx": i,
            "style": style,
            "text": p.text,
        })

# Assemble result
result = {
    "n_paragraphs": len(doc.paragraphs),
    "n_tables": len(doc.tables),
    "n_images": len(image_locations),
    "outline": outline,
    "body_order": items,
    "image_locations": image_locations,
    "metric_hits": metric_hits,
}

OUT.write_text(json.dumps(result, ensure_ascii=False, indent=2))
print(f"Wrote {OUT}")
print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}, Images: {len(image_locations)}")
print(f"Outline entries: {len(outline)}")
print(f"Metric hits: {len(metric_hits)}")
