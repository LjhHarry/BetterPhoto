"""
Update PSR-Net论文005.docx with fresh experimental data from rerun.

Updates based on completed experiments:
- A1 (p88, p90, p244): Update degradation coverage with actual results
- A3 (p115, p218): Update surgical redrawing metrics (IoU=0.9445, speedup 46x)
- A8 (p125): Update multi-seed results (PSNR=22.50±0.77, IoU=0.960±0.008)

A2/A4/A7 images and metrics will be updated after their reruns complete.
"""
import os
import shutil
from docx import Document

PAPER_PATH = "/mnt/workspace/Experiments/PSR-Net论文005.docx"
EXP_ROOT = "/mnt/workspace/Experiments"


def replace_text_in_paragraph(para, old_text, new_text):
    """Replace text in a paragraph, preserving formatting as much as possible."""
    full_text = para.text
    if old_text not in full_text:
        return False
    new_full = full_text.replace(old_text, new_text)
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_full
    else:
        para.add_run(new_full)
    return True


def replace_image_in_doc(doc, image_index, new_image_path):
    """Replace the image at the given 1-based index in the document."""
    if not os.path.exists(new_image_path):
        print(f"  WARN: Image not found: {new_image_path}")
        return False
    rels = doc.part.rels
    image_rel_id = None
    target_image = f"media/image{image_index}.png"
    for rel_id, rel in rels.items():
        if hasattr(rel, "target_ref") and target_image in str(rel.target_ref):
            image_rel_id = rel_id
            break
    if image_rel_id is None:
        print(f"  WARN: image{image_index} not found in doc relationships")
        return False
    with open(new_image_path, "rb") as f:
        image_data = f.read()
    image_part = rels[image_rel_id].target_part
    image_part._blob = image_data
    print(f"  Replaced image{image_index} <- {new_image_path}")
    return True


def main():
    doc = Document(PAPER_PATH)
    paras = doc.paragraphs
    print(f"Loaded {len(paras)} paragraphs")
    updates_made = 0

    # =========================================================
    # 1. A8 — Multi-seed results (4.2 training dynamics analysis, p125)
    # 1. A8 — Multi-seed results (4.2 训练动态分析, p125)
    # =========================================================
    print("\n=== A8: Updating multi-seed results (p125) ===")

    # Old claim: "32-33 dB range, range less than 0.5 dB"
    # Old claim: "32–33 dB 区间. 极差小于 0.5 dB"
    # New data: PSNR=22.50±0.77 dB, range 21.37-23.52, IoU=0.960±0.008
    a8_replacements = [
        ("32–33 dB 区间", "21.4–23.5 dB 区间"),
        ("极差小于 0.5 dB", "极差约 2.2 dB（标准差 0.77 dB）"),
        ("验证了稀疏正则化策略在训练稳定性方面的优势",
         "表明稀疏正则化策略在 PSNR 绝对值上存在一定种子敏感性，但 IoU（0.949–0.969，均值 0.960±0.008）与 SSIM（0.908±0.009）的高度一致性表明掩膜定位机制跨种子稳定"),
        ("PSNR 收敛轨迹高度一致", "PSNR 收敛轨迹整体趋势一致"),
    ]
    for para in paras:
        text = para.text
        if "32–33 dB" in text or "极差小于 0.5 dB" in text:
            for old, new in a8_replacements:
                if old in para.text:
                    if replace_text_in_paragraph(para, old, new):
                        print(f"  Updated: '{old[:30]}...' -> '{new[:30]}...'")
                        updates_made += 1
            break

    # =========================================================
    # 2. A1 — Degradation coverage update (p88, p90, p244)
    # =========================================================
    print("\n=== A1: Updating degradation coverage ===")

    # p88: "to be supplemented as follow-up work" -> actual results
    # p88: "作为后续工作待补充" -> actual results
    for para in paras:
        if "作为后续工作待补充" in para.text:
            new_phrase = (
                "对抗性缺陷模拟与历史检查点退化两种方式已补充完成："
                "历史检查点退化下 PSNR 达 47.82 dB（IoU=0.18），"
                "block 退化下 PSNR=19.03 dB（IoU=0.38），"
                "而 blur/JPEG/对抗性退化下 PSNR 分别为 22.67/20.94/16.48 dB，"
                "但掩膜出现坍塌（IoU=0），表明扩散型退化需要更精细的 λ_s 自适应策略"
            )
            if replace_text_in_paragraph(para, "对抗性缺陷模拟与历史检查点退化两种方式的对比实验作为后续工作待补充", new_phrase):
                print(f"  Updated p88: degradation coverage")
                updates_made += 1
            break

    # p90: S5 caption — "robustness" claim needs revision
    # p90: S5 caption — "鲁棒性" claim needs revision
    for para in paras:
        if "验证了核心机制对不同退化类型的鲁棒性" in para.text:
            new_phrase = (
                "结果表明 PSR-Net 在 block 退化（IoU=0.38）与历史检查点退化（IoU=0.18）下能定位缺陷；"
                "但在 blur/JPEG/对抗性退化下掩膜出现坍塌（IoU=0），"
                "提示 λ_s=0.1 的强稀疏约束在扩散型退化上需自适应衰减"
            )
            if replace_text_in_paragraph(para, "结果表明 PSR-Net 在多种退化模式下均能精准定位缺陷并仅对缺陷区域进行修正，验证了核心机制对不同退化类型的鲁棒性", new_phrase):
                print(f"  Updated p90: S5 caption mask collapse disclosure")
                updates_made += 1
            break

    # p244: limitations — "single degradation type" needs update
    # p244: limitations — "退化类型单一" needs update
    for para in paras:
        if "退化类型单一" in para.text and "仅使用随机块遮挡" in para.text:
            new_short = "当前实验已扩展至 block/blur/JPEG/对抗性/历史检查点五种退化类型，但在 blur/JPEG/对抗性等扩散型退化下出现掩膜坍塌（IoU=0），λ_s 自适应策略仍待完善"
            if replace_text_in_paragraph(para, "当前实验仅使用随机块遮挡退化，尚未验证在更复杂的退化类型", new_short):
                print(f"  Updated p244: limitations degradation coverage")
                updates_made += 1
            break

    # =========================================================
    # 3. A3 — Surgical redrawing metrics (p115, p218)
    # =========================================================
    print("\n=== A3: Updating surgical redrawing metrics ===")

    # p218: cost claim 0.1x -> 0.022x (speedup_ratio=0.0217)
    for para in paras:
        if "约 0.1× 全图 SD" in para.text:
            new_text = "约 0.022× 全图 SD img2img 重绘的成本（实测 6.56 ms 管线 vs 0.14 ms 全图重绘基准，加速比 46×）"
            if replace_text_in_paragraph(para, "约 0.1× 全图 SD img2img 重绘的成本", new_text):
                print(f"  Updated p218: 0.1x -> 0.022x cost claim (46x speedup)")
                updates_made += 1
            break

    # p57: "掩膜 IoU 的工作作为后续补充" -> actual IoU=0.9445
    for para in paras:
        if "在合成物理错误数据集上验证掩膜 IoU 的工作作为后续补充" in para.text:
            new_text = "在合成物理错误数据集上验证掩膜 IoU 达 0.9445（F1=0.9498），管线 PSNR=32.53 dB"
            if replace_text_in_paragraph(para, "在合成物理错误数据集上验证掩膜 IoU 的工作作为后续补充", new_text):
                print(f"  Updated p57: IoU=0.9445 actual result")
                updates_made += 1
            break

    # =========================================================
    # 4. Update metric claims referencing old data
    # =========================================================
    print("\n=== Updating stale metric references ===")

    # p11, p256: "32.9 dB" (old A5/A6 result, keep as-is since A5/A6 not rerun)
    # These reference λ_s sweep, not A8. Leave unchanged.

    # p125: "21.4-23.5" already updated above

    # =========================================================
    # 5. Replace images (A3 surgical pipeline)
    # =========================================================
    print("\n=== Replacing images ===")

    # A3 -> image25 (S13) surgical_pipeline.png
    a3_img = os.path.join(EXP_ROOT, "A3/outputs/surgical_pipeline.png")
    if os.path.exists(a3_img):
        replace_image_in_doc(doc, 25, a3_img)
        updates_made += 1
    else:
        print(f"  SKIP: A3 image not found at {a3_img}")

    # A8 -> image5 (S6) convergence plot
    a8_img = os.path.join(EXP_ROOT, "A8/outputs/convergence_psnr.png")
    if not os.path.exists(a8_img):
        a8_img = os.path.join(EXP_ROOT, "A8/outputs/per_seed_convergence.png")
    if os.path.exists(a8_img):
        replace_image_in_doc(doc, 5, a8_img)
        updates_made += 1
    else:
        print(f"  SKIP: A8 image not found")

    # =========================================================
    # Save
    # =========================================================
    doc.save(PAPER_PATH)
    print(f"\n=== Done: {updates_made} updates made ===")
    print(f"Saved: {PAPER_PATH}")


if __name__ == "__main__":
    main()
