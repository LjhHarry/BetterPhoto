"""Find table positions by scanning body XML in order, mapping each table to
its nearest preceding heading paragraph."""
from docx import Document
from docx.oxml.ns import qn

doc = Document('/mnt/workspace/Experiments/PSR-Net论文005.docx')

# Build mapping from XML element id -> paragraph index
para_by_el = {}
for i, p in enumerate(doc.paragraphs):
    para_by_el[id(p._p)] = i

table_by_el = {}
for i, t in enumerate(doc.tables):
    table_by_el[id(t._tbl)] = i

# Walk body in document order
body = doc.element.body
sequence = []  # list of (kind, idx)
last_heading = "(before any heading)"
last_heading_idx = -1

for child in body.iterchildren():
    if child.tag == qn("w:p"):
        p_idx = para_by_el.get(id(child))
        if p_idx is not None:
            p = doc.paragraphs[p_idx]
            style = p.style.name if p.style else ""
            if style.startswith("Heading") or style.startswith("标题"):
                last_heading = p.text
                last_heading_idx = p_idx
            sequence.append(("para", p_idx, last_heading, last_heading_idx))
    elif child.tag == qn("w:tbl"):
        t_idx = table_by_el.get(id(child))
        if t_idx is not None:
            sequence.append(("table", t_idx, last_heading, last_heading_idx))

# Print tables with their preceding heading
print("=== TABLES WITH SECTION CONTEXT ===")
for entry in sequence:
    if entry[0] == "table":
        t_idx = entry[1]
        heading = entry[2]
        heading_idx = entry[3]
        t = doc.tables[t_idx]
        print(f"\nTable {t_idx} (after heading '{heading}' at para {heading_idx}):")
        for r in t.rows:
            print("  | " + " | ".join(c.text[:40] for c in r.cells))

# Find paragraphs that reference specific figures (S1, S2, S5, S6, S7-S10, S13, S14, S17-S19)
print("\n=== FIGURE LABEL REFERENCES ===")
import re
fig_pattern = re.compile(r"(图|Figure|Fig\.?|S)\s*\d+|S\d{1,2}\b")
for i, p in enumerate(doc.paragraphs):
    text = p.text
    if not text.strip():
        continue
    # find figure references
    matches = list(re.finditer(r"(图\s*\d+|Figure\s*\d+|Fig\.?\s*\d+|S\d{1,2}\b)", text))
    if matches:
        # only print if it looks like a figure label
        for m in matches:
            ctx_start = max(0, m.start() - 20)
            ctx_end = min(len(text), m.end() + 40)
            print(f"  p{i}: '{m.group()}' ctx: ...{text[ctx_start:ctx_end]}...")

# Find paragraphs containing experiment IDs (A1-A9, B1-B4)
print("\n=== EXPERIMENT ID REFERENCES ===")
exp_pattern = re.compile(r"\b([AB][1-9])\b")
for i, p in enumerate(doc.paragraphs):
    text = p.text
    if not text.strip():
        continue
    matches = list(exp_pattern.finditer(text))
    if matches:
        ids = [m.group(1) for m in matches]
        # dedupe
        ids_str = ", ".join(sorted(set(ids)))
        print(f"  p{i} [{ids_str}]: {text[:120]}")
